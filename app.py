import streamlit as st
import PyPDF2
import nltk
import docx
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from PIL import Image

# Ensure nltk punkt tokenizer is downloaded
nltk.download('punkt_tab')

def summarize_text(text, num_sentences=10):  # Increased scale for summarization
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()
    summary = summarizer(parser.document, num_sentences)
    return " ".join(str(sentence) for sentence in summary)

def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        extracted_text = page.extract_text()
        if extracted_text:
            text += extracted_text + "\n"
    return text.strip() if text else "(Unable to extract text from this PDF.)"

def extract_text_from_docx(uploaded_file):
    doc = docx.Document(uploaded_file)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip() if text else "(Unable to extract text from this DOCX.)"

def main():
    st.set_page_config(page_title="🚀 AI-Powered Summarizer", layout="wide")
    
    # Custom Styling
    st.markdown(
        """
        <style>
            .big-font { font-size:25px !important; font-weight: bold; }
            .summary-box { background-color: #f0f2f6; padding: 15px; border-radius: 10px; }
            .uploaded-box { background-color: #e8f5e9; padding: 15px; border-radius: 10px; }
        </style>
        """, unsafe_allow_html=True
    )
    
    st.title("📄 AI-Powered Article Summarizer")
    st.write("💡 **Effortlessly summarize lengthy articles, PDFs, DOCX, and text files using AI-powered NLP!**")
    
    option = st.radio("📌 Choose Input Method:", ("Text Input", "Upload File"), horizontal=True)
    num_sentences = st.slider("📏 Select Number of Sentences for Summary:", 5, 25, 10)  # Increased upper limit
    text_input = ""
    
    if option == "Text Input":
        text_input = st.text_area("✍️ Enter or Paste Article Text Below:", height=300)
    elif option == "Upload File":
        uploaded_file = st.file_uploader("📂 Upload a text, PDF, or DOCX file", type=["txt", "pdf", "docx"], help="Supports .txt, .pdf, and .docx files")
        if uploaded_file is not None:
            with st.spinner("🔍 Extracting text from file..."):
                if uploaded_file.type == "application/pdf":
                    text_input = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    text_input = extract_text_from_docx(uploaded_file)
                else:
                    text_input = uploaded_file.read().decode("utf-8")
                st.markdown(f"<div class='uploaded-box'><b>📜 Extracted Content:</b></div>", unsafe_allow_html=True)
                st.text_area("", text_input, height=300)
    
    if st.button("🚀 Generate Summary", help="Click to summarize the text") and text_input:
        with st.spinner("✨ Generating summary..."):
            summary = summarize_text(text_input, num_sentences)
        st.markdown(f"<div class='summary-box'><b>✨ Summary:</b></div>", unsafe_allow_html=True)
        st.success(summary)
    # Download summary option
        st.download_button(label="📥 Download Summary as TXT", data=summary, file_name="summary.txt")
        
        doc = docx.Document()
        doc.add_paragraph(summary)
        doc.save("summary.docx")
        with open("summary.docx", "rb") as docx_file:
            st.download_button(label="📥 Download Summary as DOCX", data=docx_file, file_name="summary.docx")
        
        
        with open("summary.pdf", "rb") as pdf_file:
            st.download_button(label="📥 Download Summary as PDF", data=pdf_file, file_name="summary.pdf",mime="application/pdf")

    

if __name__ == "__main__":
    main()
